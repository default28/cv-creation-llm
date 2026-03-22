from docx import Document
from docx2pdf import convert
import subprocess

def create_docx_from_json(data, output_file="tailored_resume.docx"):
    doc = Document()

    doc.add_heading(data.get("name", ""), 0)
    doc.add_paragraph(data.get("contact", ""))

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(data.get("summary", ""))

    doc.add_heading("Skills", level=1)
    for skill in data.get("skills", []):
        doc.add_paragraph(skill, style='List Bullet')

    doc.add_heading("Experience", level=1)
    for exp in data.get("experience", []):
        doc.add_paragraph(exp)

    doc.add_heading("Projects", level=1)
    for proj in data.get("projects", []):
        doc.add_paragraph(proj)

    doc.add_heading("Education", level=1)
    doc.add_paragraph(data.get("education", ""))

    doc.save(output_file)
    return output_file


def convert_to_pdf(docx_file):
    pdf_file = docx_file.replace(".docx", ".pdf")

    try:
        convert(docx_file, pdf_file)
        return pdf_file
    except Exception as e:
        print("❌ PDF conversion failed:", e)
        return None
